import random
import time
from gtts import gTTS
from playsound import playsound
import speech_recognition as sr
import pyaudio
import os
import cv2
import re
from roboflow import Roboflow
from docx import Document
r=sr.Recognizer()

class SesliAsistan:

    def seslendirme(self, metin):
        metin_seslendirme = gTTS(text=metin, lang="tr", slow=False)
        dosya = str(random.randint(0, 10000000000)) + ".mp3"
        metin_seslendirme.save(dosya)
        playsound(dosya)
        os.remove(dosya)

    def mikrofon(self):
        with sr.Microphone() as kaynak:
            print("Sizi dinliyorum..")
            listen = r.listen(kaynak)
            ses = ""
            try:
                ses = r.recognize_google(listen, language="tr-TR")
            except sr.UnknownValueError:
                self.seslendirme("Ne dediğinizi anlayamadım")
            return ses.lower()

    def ses_karslik(self, gelen_Ses):
        if "programı kapat" in gelen_Ses:
            self.seslendirme("Program kapatılıyor")
            exit()
        elif "raporlamayı başlat" in gelen_Ses:
            self.seslendirme("Raporlama başlatılıyor")
            try:
                rf = Roboflow(api_key="")
                project = rf.workspace().project("konteyner_veri_seti")
                model = project.version("5").model

                job_id, signed_url, expire_time = model.predict_video(
                    "video5.mp4",
                    fps=5,
                    prediction_type="batch-video",
                )

                results = model.poll_until_video_results(job_id)
                print(results)
                save_list = str(results)
                rapor_dosyası = open("rapor2.doc", "w")
                rapor_dosyası.write(save_list)
                rapor_dosyası.close()

                sınıflar = ["mavi", "beyaz", "kirmizi", "kahverengi", "turuncu", "turkuvaz", "siyah", "yesil", "gri", "pembe", "sari", "tanımsız"]
                file_path1 = "rapor1.doc"
                file_path2 = "kaydedilen_rapor.doc"
                output_path = "Düzenlenmiş_Rapor.docx"

                def count_specific_words(file_path, sınıflar):
                    word_counts = {word: 0 for word in sınıflar}
                    with open(file_path, 'r', encoding='utf-8') as file:
                        for line in file:
                            words = re.findall(r'\b\w+\b', line.lower())
                            for word in words:
                                if word in word_counts:
                                    word_counts[word] += 1
                    return word_counts

                word_counts1 = count_specific_words(file_path1, sınıflar)
                word_counts2 = count_specific_words(file_path2, sınıflar)

                results_list1 = [(word, count) for word, count in word_counts1.items()]
                results_list2 = [(word, count) for word, count in word_counts2.items()]
                self.seslendirme("Buraya kadar sorunsuz çalıştı 1")

                def compare_tuples(results_list1, results_list2):
                    comparison_results = []
                    for word, expected_count in results_list2:
                        actual_count = next((count for w, count in results_list1 if w == word), 0)
                        is_match = actual_count == expected_count
                        comparison_results.append((word, f"{expected_count} kez bekleniyordu, {actual_count} kez geçti.", is_match))
                    return comparison_results

                comparison_results = compare_tuples(results_list1, results_list2)
                self.seslendirme("Buraya kadar sorunsuz çalıştı 2")

                def raporu_doldur(file_path, comparison_results):
                    doc = Document(file_path)
                    self.seslendirme("Buraya kadar sorunsuz çalıştı 3")
                    placeholders = {
                        "Şirket Adı": "………………..",
                        "Raporlama Tarihi": "..../.…/.…",
                        "Kontrol Saatleri": "..:.. / …:…",
                        "Gemi Adı": "……………….",
                        "Alıcı Adı": "………."
                    }

                    for paragraph in doc.paragraphs:
                        for key, value in placeholders.items():
                            if value in paragraph.text:
                                replacement = input(f"{key} giriniz: ")
                                paragraph.text = paragraph.text.replace(value, replacement)
                    self.seslendirme("Buraya kadar sorunsuz çalıştı 4")

                    tablo = doc.add_table(rows=1, cols=2)
                    hdr_cells = tablo.rows[0].cells
                    hdr_cells[0].text = 'Renk'
                    hdr_cells[1].text = 'Açıklama'

                    for renk, aciklama, _ in comparison_results:
                        satir_hucreler = tablo.add_row().cells
                        satir_hucreler[0].text = renk
                        satir_hucreler[1].text = aciklama

                    doc.add_paragraph('\nALICI \t\t\t\t\t\t\t\t\t\t\tGÖNDERİCİ')

                    doc.save(output_path)

                raporu_doldur("TASIMA RAPOR.docx", comparison_results)
                self.seslendirme("raporlama basarılı program kapatılıyor")
                exit()
            except Exception as e:
                self.seslendirme(f"Bir hatayla karşılaşıldı lütfen daha sonra tekrar deneyin. Hata: {str(e)}")

    def uyanma_fonksiyonu(self, gelen_Ses):
        if "hey asistan" in gelen_Ses:
            self.seslendirme("Dinliyorum...")
            ses = self.mikrofon()
            if ses:
                self.ses_karslik(ses)

asistan = SesliAsistan()

while True:
    gelen_Ses = asistan.mikrofon().lower()
    if gelen_Ses:
        print(gelen_Ses)
        asistan.uyanma_fonksiyonu(gelen_Ses)
